# api/integrations/zoho_recruit.py
"""
Zoho Recruit (minimal helpers for SME-friendly LLM routing)

Supported use cases (no Zoho API calls required):
1) Resume text extraction from small .docx files (<= 10 MB) for LLM summarization.
2) Shortlist preparation: package job criteria + multiple candidate resumes into a
   standardized, bias-minimized rubric context so the router can choose the cheapest
   capable LLM to rank candidates fairly.

Notes:
- We intentionally removed Gmail/Pipedrive/Zoho-API candidate-creation flows to keep
  this module simple and demo-ready as requested.
- All outputs are plain text for easy prompt concatenation by the router.

"""

from __future__ import annotations
import base64
import io
from typing import Any, Dict, List, Optional
try:
    from docx import Document  # used only for resume .docx reading
except Exception:
    Document = None  # allow module import even if python-docx isn't installed; handled at call time

import os
import requests

# ---- DC helpers + token cache ----
_token_cache: Dict[str, Any] = {"token": None, "ts": 0}

def _infer_region_from_base_url() -> Optional[str]:
    url = (
        os.getenv("ZOHO_RECRUIT_BASE_URL")
        or os.getenv("ZOHO_RECRUIT_API_BASE")
        or os.getenv("OHO_RECRUIT_BASE_URL")
        or ""
    ).strip().lower()
    if "zoho.eu" in url: return "eu"
    if "zoho.com.au" in url: return "au"
    if "zoho.in" in url: return "in"
    if "zoho.jp" in url: return "jp"
    if "zoho.sa" in url: return "sa"
    if "zoho.ca" in url: return "ca"
    if "zoho.com" in url: return "com"
    return (os.getenv("ZOHO_REGION") or "eu").strip().lower()

def _accounts_host_for_region(region: str) -> str:
    return {
        "eu": "accounts.zoho.eu",
        "com": "accounts.zoho.com",
        "in": "accounts.zoho.in",
        "au": "accounts.zoho.com.au",
        "jp": "accounts.zoho.jp",
        "sa": "accounts.zoho.sa",
        "ca": "accounts.zoho.ca",
    }.get(region, "accounts.zoho.eu")


def _access_token() -> str:
    """
    Prefer refresh flow (correct DC, long-lived) when available; else use static token.
    Cache tokens for ~50 minutes.
    """
    import time as _t

    # serve cached token if fresh
    if _token_cache["token"] and (_t.time() - (_token_cache["ts"] or 0) < 50 * 60):
        return _token_cache["token"]  # type: ignore[return-value]

    rt = (os.getenv("ZOHO_REFRESH_TOKEN") or "").strip()
    cid = (os.getenv("ZOHO_CLIENT_ID") or "").strip()
    cs = (os.getenv("ZOHO_CLIENT_SECRET") or "").strip()

    # If refresh triplet is provided, use it (more reliable than static access tokens)
    if rt and cid and cs:
        region = _infer_region_from_base_url()
        accounts_host = os.getenv("ZOHO_ACCOUNTS_HOST", _accounts_host_for_region(region or "eu"))
        url = f"https://{accounts_host}/oauth/v2/token"
        r = requests.post(
            url,
            data={
                "grant_type": "refresh_token",
                "refresh_token": rt,
                "client_id": cid,
                "client_secret": cs,
            },
            timeout=30,
        )
        r.raise_for_status()
        js = r.json()
        tok = js.get("access_token")
        if not tok:
            raise RuntimeError("Zoho refresh returned no access_token")
        _token_cache["token"] = tok
        _token_cache["ts"] = _t.time()
        return tok

    # Fallback: static access token from env
    tok = os.getenv("ZOHO_ACCESS_TOKEN")
    if not tok:
        raise RuntimeError("Missing ZOHO_ACCESS_TOKEN or (ZOHO_REFRESH_TOKEN + ZOHO_CLIENT_ID + ZOHO_CLIENT_SECRET)")
    _token_cache["token"] = tok
    _token_cache["ts"] = _t.time()
    return tok




# --------- Gmail wrapper (optional) ----------
try:
    from api.integrations.google_ws import gmail_fetch_newest_thread
except Exception:
    gmail_fetch_newest_thread = None


# --------- Resume prefetch (.docx <=10MB) ----------
def prefetch_resume_b64(b64: str, filename: str = "resume.docx", **_) -> str:
    if Document is None:
        return "Server is missing python-docx; please add it to requirements and rebuild."
    if not b64:
        return "No resume provided."
    try:
        raw = base64.b64decode(b64, validate=True)
    except Exception:
        return "Invalid base64."
    size_mb = len(raw) / (1024 * 1024)
    if size_mb > 10.0:
        return "Resume too large (>10MB)."
    if not (filename or "").lower().endswith(".docx"):
        return "Unsupported file type (only .docx)."
    try:
        fileobj = io.BytesIO(raw)
        doc = Document(fileobj)
        chunks: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                chunks.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join([c for c in cells if c])
                if line:
                    chunks.append(line)
        text = "\n".join(chunks).strip()
        if not text:
            return "Resume extracted but contains no readable text."
        return (text[:20000] + "\n[...trimmed...]") if len(text) > 20000 else text
    except Exception as e:
        return f"Failed to read .docx: {e}"


def shortlist_prefetch(candidates: List[Dict[str, Any]], job_criteria: Dict[str, Any], **_) -> str:
    """
    Prepare a standardized, bias-minimized context for LLM shortlisting.
    Inputs:
      candidates: [{"name": str, "resume_b64": str, "filename": "x.docx"}, ...]
      job_criteria: {
        "title": str (optional),
        "must_have": [str, ...],
        "nice_to_have": [str, ...],
        "min_years_experience": int (optional),
        "required_qualifications": [str, ...],
        "location": [str, ...] (optional),
        "keywords": [str, ...] (optional),
        "other": str (optional)
      }
    Output: Plain text with:
      - A bias-minimized rubric
      - The normalized job criteria
      - Per-candidate packs with extracted resume text (truncated)
    """
    # ---- 1) Bias-minimized scoring rubric (keep consistent across tasks) ----
    rubric = [
        "### Standardized, Bias-Minimized Rubric (0–5 each, integers only)",
        "1) Skills Match: how well tangible skills match MUST-HAVEs; partial credit for NICE-TO-HAVEs.",
        "2) Relevant Experience: years and depth in matching roles/tech/domains (evidence-based).",
        "3) Qualifications: required certs/degrees/licenses; partial for equivalents.",
        "4) Achievements & Impact: quantified outcomes, scope, leadership, publications, awards.",
        "5) Role Fit Signals: availability, location overlap, stability, communication clarity.",
        "",
        "Rules to reduce bias: Do NOT consider name, gender, age, nationality, photo, address, or school prestige.",
        "Rely only on verifiable evidence in the resume text. If unclear, mark as 'Insufficient Evidence'.",
    ]

    # ---- 2) Normalize job criteria ----
    def _fmt_list(key: str) -> str:
        vals = job_criteria.get(key) or []
        if isinstance(vals, str):
            vals = [vals]
        vals = [str(v).strip() for v in vals if str(v).strip()]
        return f"- {key.replace('_',' ').title()}: " + (", ".join(vals) if vals else "None specified")

    jc_lines: List[str] = []
    title = str(job_criteria.get("title") or "").strip()
    if title:
        jc_lines.append(f"**Role Title**: {title}")
    jc_lines.append(_fmt_list("must_have"))
    jc_lines.append(_fmt_list("nice_to_have"))
    min_yrs = job_criteria.get("min_years_experience")
    if isinstance(min_yrs, (int, float)):
        jc_lines.append(f"- Minimum Years Experience: {int(min_yrs)}")
    jc_lines.append(_fmt_list("required_qualifications"))
    jc_lines.append(_fmt_list("location"))
    jc_lines.append(_fmt_list("keywords"))
    other = str(job_criteria.get("other") or "").strip()
    if other:
        jc_lines.append(f"- Other: {other}")

    # ---- 3) Build candidate packs ----
    packs: List[str] = []
    MAX_PER_RESUME = 12000  # keep LLM cost low
    for idx, c in enumerate(candidates, 1):
        name = str(c.get("name") or f"Candidate {idx}").strip()
        b64 = c.get("resume_b64") or ""
        fn = c.get("filename") or "resume.docx"
        extracted = prefetch_resume_b64(b64, fn)
        if isinstance(extracted, str) and len(extracted) > MAX_PER_RESUME:
            extracted = extracted[:MAX_PER_RESUME] + "\n[...trimmed for cost...]"
        packs.append(
            f"## Candidate {idx}: {name}\n"
            f"ResumeText:\n{extracted}\n"
        )

    # ---- 4) Combine into final context ----
    out = []
    out.extend(rubric)
    out.append("\n### Job Criteria (Normalized)")
    out.append("\n".join(jc_lines))
    out.append("\n### Candidates")
    out.append("\n\n".join(packs))
    return "\n".join(out).strip()


# ---- Optional Zoho Recruit fetchers (Candidates -> Attachments -> download .docx) ----
# Requires: ZOHO_ACCESS_TOKEN (simple demo) and either ZOHO_RECRUIT_BASE_URL or ZOHO_REGION
def _zr_base() -> str:
    # Prefer explicit base URL; accept both keys + a common typo for resilience.
    override = (
        os.getenv("ZOHO_RECRUIT_BASE_URL")
        or os.getenv("ZOHO_RECRUIT_API_BASE")
        or os.getenv("OHO_RECRUIT_BASE_URL")  # typo-tolerant
    )
    if override:
        return override.rstrip("/")

    # Fallback: region-based host mapping
    region = (os.getenv("ZOHO_REGION", "eu") or "eu").strip().lower()
    host = {
        "eu": "recruit.zoho.eu",
        "com": "recruit.zoho.com",
        "in": "recruit.zoho.in",
        "au": "recruit.zoho.com.au",
        "jp": "recruit.zoho.jp",
        "sa": "recruit.zoho.sa",
        "ca": "recruit.zoho.ca",
    }.get(region, "recruit.zoho.eu")
    return f"https://{host}/recruit/v2"


def _zr_headers() -> dict:
    h = {"Authorization": f"Zoho-oauthtoken {_access_token()}"}
    org = os.getenv("ZOHO_ORG_ID")
    if org:
        h["X-RECRUIT-ORG"] = org
    return h

def _zr_get(path: str, params: dict | None = None, stream: bool = False):
    import json as _json
    url = f"{_zr_base()}{path}"

    def _do() -> requests.Response:
        return requests.get(url, headers=_zr_headers(), params=params, timeout=30, stream=stream)

    r = _do()
    if r.status_code == 401:
        # If refresh flow is configured, try to refresh and retry once
        if os.getenv("ZOHO_REFRESH_TOKEN") and os.getenv("ZOHO_CLIENT_ID") and os.getenv("ZOHO_CLIENT_SECRET"):
            # Clear cache and re-acquire
            _token_cache["token"] = None
            _token_cache["ts"] = 0
            _ = _access_token()
            r = _do()

    if r.status_code >= 400:
        try:
            j = r.json()
            msg = _json.dumps(j)[:300]
        except Exception:
            msg = r.text[:300]
        raise RuntimeError(f"Zoho GET {path} failed {r.status_code}: {msg}")
    return r


def _zr_get_json(path: str, params: dict | None = None) -> dict:
    r = _zr_get(path, params=params, stream=False)
    try:
        return r.json()
    except Exception:
        return {}

def fetch_resume_text_from_zoho(candidate_id: str) -> str:
    """
    Find first .docx attachment for the Candidate, download it,
    and return extracted text (<=10MB, parsed via existing prefetch).
    """
    if not candidate_id:
        return "No candidate_id provided."
    # 1) list attachments
    j = _zr_get_json(f"/Candidates/{candidate_id}/Attachments")
    items = (j.get("data") or []) if isinstance(j, dict) else []
    if not items:
        return "No attachments for candidate."
    # 2) pick a .docx (smallest first)
    def _pick_key(it):
        name = (it.get("File_Name") or it.get("file_name") or "").lower()
        size = int(it.get("Size") or it.get("size") or 0)
        return (0 if name.endswith(".docx") else 1, size or 10**9)
    items_sorted = sorted(items, key=_pick_key)
    chosen = None
    for it in items_sorted:
        name = (it.get("File_Name") or it.get("file_name") or "")
        if name.lower().endswith(".docx"):
            chosen = it
            break
    if not chosen:
        return "No .docx attachments for candidate."
    att_id = chosen.get("id") or chosen.get("ID") or chosen.get("Attachment_Id")
    fname = chosen.get("File_Name") or chosen.get("file_name") or "resume.docx"
    if not att_id:
        return "Attachment id not found."
    # 3) download bytes
    r = _zr_get(f"/Candidates/{candidate_id}/Attachments/{att_id}", stream=True)
    data = r.content
    if len(data) > 10 * 1024 * 1024:
        return "Resume too large (>10MB)."
    # 4) reuse the existing .docx parser by base64-encoding
    import base64 as _b64
    return prefetch_resume_b64(_b64.b64encode(data).decode("ascii"), fname)

def shortlist_prefetch_from_zoho(candidate_ids: List[str], job_criteria: Dict[str, Any], **_) -> str:
    """
    Build the same shortlist context as shortlist_prefetch, but fetch resumes by Candidate IDs.
    """
    candidates: List[Dict[str, Any]] = []
    for idx, cid in enumerate(candidate_ids, 1):
        # try to get candidate name (best-effort)
        name = f"Candidate {idx}"
        try:
            rec = _zr_get_json(f"/Candidates/{cid}")  # single record
            data = rec.get("data") or []
            if isinstance(data, list) and data:
                row = data[0]
                first = (row.get("First_Name") or "").strip()
                last = (row.get("Last_Name") or "").strip()
                nm = f"{first} {last}".strip() or row.get("Email") or name
                name = nm
        except Exception:
            pass
        extracted = fetch_resume_text_from_zoho(cid)
        candidates.append({"name": name, "resume_b64": "", "filename": "from_zoho.docx", "extracted": extracted})

    # Convert 'extracted' to the expected format by shortlist_prefetch
    # (it calls prefetch_resume_b64 internally, so we pass the text directly)
    # We'll adapt: if 'extracted' already plain text, use as-is.
    # Rebuild packs using the same shape:
    rubric_context = shortlist_prefetch(
        candidates=[{"name": c["name"], "resume_b64": "", "filename": "from_zoho.docx"} for c in candidates],
        job_criteria=job_criteria
    )
    # Replace the auto-extracted placeholder with our fetched text
    # (quick substitution—each "Candidate N: Name\nResumeText:\n" anchor is unique enough for demo use)
    out_lines = rubric_context.split("### Candidates", 1)
    head = out_lines[0]
    body = out_lines[1] if len(out_lines) > 1 else ""
    built = ["### Candidates"]
    for idx, c in enumerate(candidates, 1):
        built.append(f"\n## Candidate {idx}: {c['name']}\nResumeText:\n{c['extracted']}\n")
    return (head + "\n" + "\n".join(built)).strip()

def resume_summarize_prefetch_from_zoho(candidate_id: str, **_) -> str:
    """
    Fetch the candidate's first .docx attachment from Zoho and return plain text for LLM summarization.
    """
    return fetch_resume_text_from_zoho(candidate_id)
