# api/integrations/ms365.py
# Microsoft 365 helpers used by the router.
# FIX: use ONLY resource scopes in acquire_token_silent (no reserved scopes).

from __future__ import annotations

import os
import json
import datetime as dt
from typing import List, Optional

import httpx
import msal

import io
from docx import Document
from docx.shared import Pt


# ------------------------------------------------------------------------------
# Configuration
# ------------------------------------------------------------------------------

GRAPH = "https://graph.microsoft.com/v1.0"
AUTHORITY = f"https://login.microsoftonline.com/{os.getenv('MS_TENANT_ID','consumers')}"
CLIENT_ID = os.getenv("MS_CLIENT_ID", "")

# Use ONLY resource scopes for silent refresh (no reserved scopes here)
RESOURCE_SCOPES = [
    "User.Read",
    "Mail.ReadWrite",
    "Mail.Send",
    "Calendars.ReadWrite",
    "Files.ReadWrite",
]


# ------------------------------------------------------------------------------
# Token cache path (robust)
# ------------------------------------------------------------------------------

def _candidate_token_paths() -> List[str]:
    """Return a prioritized list of token-cache paths to try."""
    override = os.getenv("MS_TOKEN_CACHE_PATH")
    if override:
        return [override]
    return [
        "/app/secrets/ms_token_cache.json",      # ← bind mount target
        "/app/api/secrets/ms_token_cache.json",
        "/app/app/secrets/ms_token_cache.json",
        "/secrets/ms_token_cache.json",
        os.path.abspath("./secrets/ms_token_cache.json"),
    ]


def _resolve_token_cache_path() -> str:
    """Pick the first existing token cache path or raise with detailed info."""
    tried = []
    for p in _candidate_token_paths():
        tried.append(p)
        if os.path.exists(p):
            return p
    raise RuntimeError(
        "Token cache not found. Tried:\n  - " + "\n  - ".join(tried) +
        "\nFix: ensure docker-compose binds ./secrets to /app/secrets and re-run:\n"
        "  docker compose exec api python scripts/ms_auth.py"
    )


def _load_token() -> str:
    """
    Load a valid Graph access token from the MSAL device-code cache.
    Uses ONLY resource scopes for silent refresh.
    """
    cache_path = _resolve_token_cache_path()

    cache = msal.SerializableTokenCache()
    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            cache.deserialize(f.read())
    except Exception as e:
        raise RuntimeError(f"Token cache unreadable at {cache_path}: {e}")

    app = msal.PublicClientApplication(CLIENT_ID, authority=AUTHORITY, token_cache=cache)
    accounts = app.get_accounts()
    if not accounts:
        raise RuntimeError("No Microsoft account found in cache. Re-run ms_auth.py.")

    result = app.acquire_token_silent(RESOURCE_SCOPES, account=accounts[0])
    if not result or "access_token" not in result:
        raise RuntimeError("Cannot refresh token silently. Re-run ms_auth.py.")
    return result["access_token"]


# ------------------------------------------------------------------------------
# Small utilities
# ------------------------------------------------------------------------------

def _strip_text(s: Optional[str]) -> str:
    if not s:
        return ""
    s = s.replace("\r", " ").replace("\n", " ").replace("\t", " ")
    while "  " in s:
        s = s.replace("  ", " ")
    return s.strip()


def _iso_ago(days: int) -> str:
    t = dt.datetime.utcnow() - dt.timedelta(days=days)
    return t.strftime("%Y-%m-%dT%H:%M:%SZ")


# ------------------------------------------------------------------------------
# TRIAGE CONTEXT (used by integration kind: ms.mail_triage)
# ------------------------------------------------------------------------------

def build_inbox_triage_context(n: int = 3, lookback_days: int = 30) -> str:
    """
    Fetch the latest N messages from Inbox, plus a minimal sender-history count,
    and return a compact, LLM-friendly text block.
    """
    token = _load_token()
    headers = {"Authorization": f"Bearer {token}"}

    params = {
        "$top": str(n),
        "$orderby": "receivedDateTime desc",
        "$select": "id,subject,from,receivedDateTime,bodyPreview",
    }
    try:
        with httpx.Client(timeout=30) as client:
            r = client.get(f"{GRAPH}/me/mailFolders/Inbox/messages", headers=headers, params=params)
            r.raise_for_status()
            data = r.json()
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"Graph error (Inbox fetch): {e.response.status_code} {e.response.text[:200]}")
    except Exception as e:
        raise RuntimeError(f"Graph error (Inbox fetch): {e}")

    msgs: List[dict] = data.get("value", [])[:n]
    lines: List[str] = []
    since = _iso_ago(lookback_days)

    for i, m in enumerate(msgs, start=1):
        subj = _strip_text(m.get("subject", ""))
        preview = _strip_text(m.get("bodyPreview", ""))
        received = m.get("receivedDateTime", "")
        sender = ((m.get("from") or {}).get("emailAddress") or {}).get("address", "")

        # Minimal sender history for lookback window
        hist_count = 0
        if sender:
            hist_filter = f"from/emailAddress/address eq '{sender}' and receivedDateTime ge {since}"
            hist_params = {"$select": "id", "$filter": hist_filter, "$top": "50"}
            try:
                with httpx.Client(timeout=30) as client:
                    rr = client.get(f"{GRAPH}/me/messages", headers=headers, params=hist_params)
                    if rr.status_code == 200:
                        hist_count = len(rr.json().get("value", []))
            except Exception:
                pass

        lines.append(
            f"[{i}] From: {sender} | Received: {received}\n"
            f"Subject: {subj}\n"
            f"Preview: {preview}\n"
            f"Sender history (since {since}): count={hist_count}"
        )

    return "\n".join(lines) if lines else "No recent messages found in Inbox."


# ------------------------------------------------------------------------------
# Compatible signatures used by existing routes
# ------------------------------------------------------------------------------

def calendar_create(subject: str, body: str, start: Optional[str], end: Optional[str], tz: str = "Europe/Warsaw") -> str:
    token = _load_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

    if not start or not end:
        tomorrow = dt.datetime.now(dt.timezone.utc) + dt.timedelta(days=1)
        start_dt = dt.datetime(tomorrow.year, tomorrow.month, tomorrow.day, 8, 0, tzinfo=dt.timezone.utc)  # ~10:00 CET
        end_dt = start_dt + dt.timedelta(minutes=30)
        start = start_dt.isoformat().replace("+00:00", "Z")
        end = end_dt.isoformat().replace("+00:00", "Z")

    payload = {
        "subject": subject,
        "body": {"contentType": "text", "content": body or ""},
        "start": {"dateTime": start, "timeZone": "UTC"},
        "end": {"dateTime": end, "timeZone": "UTC"},
    }

    with httpx.Client(timeout=30) as client:
        r = client.post(f"{GRAPH}/me/events", headers=headers, json=payload)
        r.raise_for_status()
        ev = r.json()
        return ev.get("webLink", "")


def mail_draft_reply_latest(body_text: str) -> str:
    token = _load_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

    params = {"$top": "1", "$orderby": "receivedDateTime desc", "$select": "id,webLink"}
    with httpx.Client(timeout=30) as client:
        r = client.get(f"{GRAPH}/me/mailFolders/Inbox/messages", headers=headers, params=params)
        r.raise_for_status()
        vals = r.json().get("value", [])
        if not vals:
            raise RuntimeError("No messages to reply to.")
        msg_id = vals[0]["id"]

        r2 = client.post(f"{GRAPH}/me/messages/{msg_id}/createReply", headers=headers)
        r2.raise_for_status()
        draft = r2.json()
        draft_id = draft["id"]

        patch = {"body": {"contentType": "text", "content": body_text}}
        r3 = client.patch(f"{GRAPH}/me/messages/{draft_id}", headers=headers, json=patch)
        r3.raise_for_status()

        r4 = client.get(f"{GRAPH}/me/messages/{draft_id}?$select=webLink", headers=headers)
        if r4.status_code == 200:
            return r4.json().get("webLink", draft_id)
        return draft_id


# add at top with other imports
import io
from docx import Document
from docx.shared import Pt

def word_upsert_docx(text: str, file_name: str, folder: Optional[str] = None) -> str:
    """
    Build a valid .docx in memory (python-docx) and upload it via Graph simple upload.
    This avoids the 'cannot open online' error caused by pushing raw UTF-8 text.
    Returns a Word Online webUrl.
    """
    token = _load_token()
    headers = {
        "Authorization": f"Bearer {token}",
        # IMPORTANT: send correct DOCX content type
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    # 1) Build a real DOCX (Office Open XML) in memory
    doc = Document()
    # (light styling to ensure a proper document body)
    try:
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)
    except Exception:
        pass

    # Title + body
    doc.add_heading("Generated by LLM Router", level=1)
    doc.add_paragraph("")
    for line in (text or "").splitlines():
        doc.add_paragraph(line if line.strip() else "")

    bio = io.BytesIO()
    doc.save(bio)
    data = bio.getvalue()

    # sanity guard: a valid .docx shouldn't be tiny
    if len(data) < 800:
        raise RuntimeError("DOCX generation failed (document too small).")

    # 2) Upload to OneDrive (create or replace)
    if folder:
        drive_path = f"/me/drive/root:/{folder}/{file_name}:/content"
    else:
        drive_path = f"/me/drive/root:/{file_name}:/content"

    with httpx.Client(timeout=60) as client:
        r = client.put(f"{GRAPH}{drive_path}", headers=headers, content=data)
        r.raise_for_status()
        item = r.json()
        item_id = item.get("id")
        if not item_id:
            return ""
        r2 = client.get(f"{GRAPH}/me/drive/items/{item_id}?select=webUrl", headers={"Authorization": f"Bearer {token}"})
        if r2.status_code == 200:
            return r2.json().get("webUrl", "")
        return ""
